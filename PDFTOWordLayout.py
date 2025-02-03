import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import os
import io
from PIL import Image

def extract_images(pdf_file, word_file):
    try:
        print("Starting the conversion process...")  # Debugging start message
        
        # Open the PDF with PyMuPDF
        doc = fitz.open(pdf_file)
        docx = Document()

        # Determine the directory to save temporary images (same directory as PDF or Word file)
        output_dir = os.path.dirname(word_file)  # Use the Word file's directory
        
        # Ensure directory exists
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Iterate over each page in the PDF
        for page_num in range(len(doc)):
            print(f"Processing page {page_num + 1}...")  # Debugging page number
            
            page = doc.load_page(page_num)
            
            # Create a table to simulate the PDF layout (1 row with 2 columns)
            table = docx.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            # Extract text from the page
            text = page.get_text("text")
            # Place the text in the first column of the table
            table.cell(0, 0).text = text
            
            # Extract images on the page
            img_list = page.get_images(full=True)
            for img_index, img in enumerate(img_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                # Convert image bytes to Image object (PIL)
                image = Image.open(io.BytesIO(image_bytes))
                
                # Save the image temporarily to insert into Word
                image_path = os.path.join(output_dir, f"temp_image_{page_num}_{img_index}.png")
                image.save(image_path)

                # Insert the image into the second column of the table (top-left corner of Word)
                cell = table.cell(0, 1)
                cell.paragraphs[0].clear()  # Clear existing text in cell
                cell.paragraphs[0].add_run().add_picture(image_path, width=Inches(2.0))  # Adjust the size as needed
                
                # Optionally remove the temporary image file
                os.remove(image_path)

        # Save the Word document
        docx.save(word_file)
        print(f"Conversion completed. The Word file is saved as {word_file}")
    
    except Exception as e:
        # If an error occurs, print the error message and log it
        print(f"ERROR: An error occurred during conversion: {e}")
        
        # Create the error log file in the same directory as the output Word file
        error_log_path = os.path.join(output_dir, "error_log.txt")
        with open(error_log_path, "w") as log_file:
            log_file.write(f"Error occurred: {str(e)}\n")
        
        print(f"An error occurred. Check the error log in: {error_log_path}")
        
        # Pauser to allow you to check the error log
        input("Press Enter to exit and check the error log...")

# Pauser to prevent automatic execution when you run the script
input("Press Enter to start the conversion process...")

# Example usage
pdf_file = r'C:\path\to\your\input.pdf'  # Replace with your input PDF file path
word_file = r'C:\path\to\your\output_with_images.docx'  # Replace with your desired output Word file path


# Call the function to perform conversion
extract_images(pdf_file, word_file)

# Pauser to keep the program open and allow you to read any output or error messages
print("Conversion process completed or encountered an error.")
input("Press Enter to close the program after checking results...")
